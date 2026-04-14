from __future__ import annotations

import base64
from dataclasses import dataclass
import io
import json
import time
from pathlib import Path
from urllib.parse import urlencode, urlparse
from urllib.request import Request, urlopen

from school_guardian.config import Settings
from school_guardian.domain import ClassroomTask, TaskMaterial


_SUPPORTED_URL_SUFFIXES = {".docx", ".pdf", ".png", ".jpg", ".jpeg", ".gif", ".webp"}
_SUPPORTED_IMAGE_MIME_PREFIX = "image/"
_SUPPORTED_BINARY_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
_GOOGLE_EXPORT_MIME_TYPES = {
    "application/vnd.google-apps.document": "application/pdf",
    "application/vnd.google-apps.presentation": "application/pdf",
}


@dataclass(frozen=True)
class MaterialBlob:
    filename: str
    mime_type: str
    data: bytes


@dataclass(frozen=True)
class MaterialExtraction:
    text: str | None
    source: str | None


class MaterialDownloadService:
    def __init__(self, settings: Settings) -> None:
        self.settings = settings

    def download_supported_materials(self, tasks: list[ClassroomTask]) -> list[Path]:
        downloaded: list[Path] = []
        for task in tasks:
            task_dir = self.settings.download_dir / _safe_name(task.course_name) / _safe_name(task.title)
            task_dir.mkdir(parents=True, exist_ok=True)
            for material in task.materials:
                downloaded_path = self._download_material(material, task_dir)
                if downloaded_path is not None:
                    downloaded.append(downloaded_path)
        return downloaded

    def _download_material(self, material: TaskMaterial, task_dir: Path) -> Path | None:
        if material.material_type == "link" and material.url:
            suffix = Path(urlparse(material.url).path).suffix.lower()
            if suffix in {".pdf", ".docx"}:
                return self._download_url(material.url, task_dir / f"{_safe_name(material.title)}{suffix}")
            return None

        if material.material_type == "drive_file" and material.drive_file_id:
            return self._download_google_drive_material(material, task_dir)

        return None

    def _download_google_drive_material(self, material: TaskMaterial, task_dir: Path) -> Path | None:
        try:
            from google.oauth2.credentials import Credentials
            from googleapiclient.discovery import build
            from googleapiclient.http import MediaIoBaseDownload
        except ImportError as exc:
            raise RuntimeError(
                "Descargar archivos de Drive requiere las dependencias opcionales de Google."
            ) from exc

        if not self.settings.google_token_path.exists():
            raise RuntimeError("No existe el token OAuth para descargar materiales de Drive.")

        creds = Credentials.from_authorized_user_file(
            str(self.settings.google_token_path), self.settings.google_scopes
        )
        service = build("drive", "v3", credentials=creds)
        metadata = (
            service.files()
            .get(fileId=material.drive_file_id, fields="id,name,mimeType")
            .execute()
        )

        mime_type = metadata["mimeType"]
        filename = metadata["name"]
        if mime_type == "application/pdf":
            request = service.files().get_media(fileId=material.drive_file_id)
            destination = task_dir / f"{_safe_name(filename)}.pdf"
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            request = service.files().get_media(fileId=material.drive_file_id)
            destination = task_dir / f"{_safe_name(filename)}.docx"
        elif mime_type == "application/vnd.google-apps.document":
            request = service.files().export_media(
                fileId=material.drive_file_id,
                mimeType="application/pdf",
            )
            destination = task_dir / f"{_safe_name(filename)}.pdf"
        else:
            return None

        with destination.open("wb") as handle:
            downloader = MediaIoBaseDownload(handle, request)
            done = False
            while not done:
                _, done = downloader.next_chunk()
        return destination

    def _download_url(self, url: str, destination: Path) -> Path:
        with urlopen(url) as response:
            destination.write_bytes(response.read())
        return destination


def dump_task_material_manifest(tasks: list[ClassroomTask], destination: Path) -> Path:
    payload = []
    for task in tasks:
        payload.append(
            {
                "external_id": task.external_id,
                "course_name": task.course_name,
                "title": task.title,
                "materials": [
                    {
                        "material_id": material.material_id,
                        "title": material.title,
                        "material_type": material.material_type,
                        "url": material.url,
                        "drive_file_id": material.drive_file_id,
                    }
                    for material in task.materials
                ],
            }
        )
    destination.write_text(json.dumps(payload, indent=2, ensure_ascii=False))
    return destination


def extract_text_from_material(material: TaskMaterial, settings: Settings) -> str | None:
    extracted = extract_text_with_source_from_material(material, settings)
    return extracted.text


def extract_text_with_source_from_material(
    material: TaskMaterial, settings: Settings
) -> MaterialExtraction:
    """Download and extract plain text from a material. Returns None if unsupported or failed."""
    if settings.azure_document_intelligence_endpoint and settings.azure_document_intelligence_key:
        blob = _resolve_material_blob(material, settings)
        if blob is not None:
            extracted = _analyze_with_document_intelligence(blob, settings)
            if extracted:
                return MaterialExtraction(text=extracted, source="azure_document_intelligence")
        else:
            print(
                json.dumps(
                    {
                        "event": "document_intelligence_skipped",
                        "reason": "unsupported_material",
                        "material_type": material.material_type,
                        "material_id": material.material_id,
                        "title": material.title,
                    },
                    ensure_ascii=False,
                )
            )
    else:
        print(
            json.dumps(
                {
                    "event": "document_intelligence_skipped",
                    "reason": "not_configured",
                    "material_type": material.material_type,
                    "material_id": material.material_id,
                    "title": material.title,
                },
                ensure_ascii=False,
            )
        )

    if material.material_type == "drive_file" and material.drive_file_id:
        return MaterialExtraction(
            text=_extract_from_drive(material, settings),
            source="drive_fallback",
        )
    if material.material_type == "link" and material.url:
        return MaterialExtraction(
            text=_extract_from_url(material.url),
            source="url_fallback",
        )
    return MaterialExtraction(text=None, source=None)


def _extract_from_drive(material: TaskMaterial, settings: Settings) -> str | None:
    try:
        from google.oauth2.credentials import Credentials
        from googleapiclient.discovery import build
        from googleapiclient.http import MediaIoBaseDownload
    except ImportError:
        return None

    if not settings.google_token_path.exists():
        return None

    try:
        creds = Credentials.from_authorized_user_file(str(settings.google_token_path), settings.google_scopes)
        service = build("drive", "v3", credentials=creds)

        mime_type = material.mime_type
        if not mime_type:
            metadata = service.files().get(fileId=material.drive_file_id, fields="mimeType").execute()
            mime_type = metadata.get("mimeType", "")

        # Google Workspace docs/slides → export as plain text
        if mime_type in {
            "application/vnd.google-apps.document",
            "application/vnd.google-apps.presentation",
        }:
            raw = service.files().export_media(fileId=material.drive_file_id, mimeType="text/plain").execute()
            return raw.decode("utf-8", errors="replace") if isinstance(raw, bytes) else str(raw)

        # Native DOCX in Drive
        if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            buf = io.BytesIO()
            downloader = MediaIoBaseDownload(buf, service.files().get_media(fileId=material.drive_file_id))
            done = False
            while not done:
                _, done = downloader.next_chunk()
            buf.seek(0)
            return _extract_docx(buf)

        # Native PDF in Drive
        if mime_type == "application/pdf":
            buf = io.BytesIO()
            downloader = MediaIoBaseDownload(buf, service.files().get_media(fileId=material.drive_file_id))
            done = False
            while not done:
                _, done = downloader.next_chunk()
            buf.seek(0)
            return _extract_pdf(buf)

    except Exception as exc:
        print(json.dumps({"event": "extract_drive_error", "drive_file_id": material.drive_file_id, "error": str(exc)}, ensure_ascii=False))
        return None

    print(json.dumps({"event": "extract_drive_unsupported_mime", "drive_file_id": material.drive_file_id, "mime_type": mime_type}, ensure_ascii=False))
    return None


def _download_google_drive_material_bytes(material: TaskMaterial, settings: Settings) -> bytes | None:
    """Download raw bytes of a Drive file (DOCX only). Returns None on failure."""
    try:
        from google.oauth2.credentials import Credentials
        from googleapiclient.discovery import build
        from googleapiclient.http import MediaIoBaseDownload
    except ImportError:
        return None

    if not settings.google_token_path.exists():
        return None

    try:
        creds = Credentials.from_authorized_user_file(str(settings.google_token_path), settings.google_scopes)
        service = build("drive", "v3", credentials=creds)

        mime_type = material.mime_type
        if not mime_type:
            metadata = service.files().get(fileId=material.drive_file_id, fields="mimeType").execute()
            mime_type = metadata.get("mimeType", "")

        if mime_type != "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return None

        buf = io.BytesIO()
        downloader = MediaIoBaseDownload(buf, service.files().get_media(fileId=material.drive_file_id))
        done = False
        while not done:
            _, done = downloader.next_chunk()
        return buf.getvalue()
    except Exception as exc:
        print(json.dumps({"event": "download_drive_error", "drive_file_id": material.drive_file_id, "error": str(exc)}, ensure_ascii=False))
        return None


def _extract_from_url(url: str) -> str | None:
    suffix = Path(urlparse(url).path).suffix.lower()
    if suffix not in {".docx", ".pdf"}:
        return None
    try:
        with urlopen(url, timeout=15) as response:
            content = response.read()
    except Exception:
        return None
    if suffix == ".docx":
        return _extract_docx(io.BytesIO(content))
    return _extract_pdf(io.BytesIO(content))


def _extract_docx(buf: io.BytesIO) -> str | None:
    try:
        import zipfile

        from docx import Document

        doc = Document(buf)
        text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text += "\n" + row_text

        if text.strip():
            return text.strip()

        return None
    except Exception:
        return None


def extract_images_from_docx(buf: io.BytesIO) -> list[tuple[str, bytes]]:
    """Return list of (filename, raw_bytes) for images embedded in a DOCX file."""
    import zipfile

    images: list[tuple[str, bytes]] = []
    try:
        buf.seek(0)
        with zipfile.ZipFile(buf) as zf:
            for name in zf.namelist():
                if name.startswith("word/media/"):
                    filename = name.split("/")[-1]
                    images.append((filename, zf.read(name)))
    except Exception:
        pass
    return images


def _extract_pdf(buf: io.BytesIO) -> str | None:
    try:
        from pypdf import PdfReader
        reader = PdfReader(buf)
        pages = [page.extract_text() for page in reader.pages]
        text = "\n".join(t for t in pages if t)
        return text or None
    except Exception:
        return None


def analyze_images_with_vision(
    images: list[tuple[str, bytes]],
    api_key: str,
    base_url: str,
    deployment: str,
) -> str | None:
    """Send images to an Azure OpenAI vision model and return the text analysis."""
    if not images:
        return None

    # Build the chat completions URL.
    # New-style base URLs end with /openai/v1/ — append chat/completions directly.
    # Classic Azure URLs point to the resource root — include deployment path.
    base = base_url.rstrip("/")
    if "/v1" in base or "deployments/" in base:
        url = f"{base}/chat/completions"
    else:
        url = f"{base}/openai/deployments/{deployment}/chat/completions?api-version=2024-10-21"

    content: list[dict] = [
        {
            "type": "text",
            "text": (
                "Esta es una imagen de material de una tarea escolar. "
                "Describí en detalle todo el contenido que ves: texto, tablas, diagramas, etc. "
                "Si hay texto, transcribilo íntegro. Si hay una tabla, reproducí su contenido."
            ),
        }
    ]
    for filename, img_bytes in images:
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "png"
        mime = f"image/{ext}" if ext in {"png", "jpg", "jpeg", "gif", "webp"} else "image/png"
        b64 = base64.b64encode(img_bytes).decode()
        content.append({"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}})

    payload = json.dumps({
        "model": deployment,
        "messages": [{"role": "user", "content": content}],
        "max_tokens": 2000,
    }).encode()

    try:
        req = Request(
            url,
            data=payload,
            headers={"api-key": api_key, "Content-Type": "application/json"},
        )
        with urlopen(req, timeout=60) as resp:
            data = json.loads(resp.read())
        return data["choices"][0]["message"]["content"]
    except Exception as exc:
        print(json.dumps({"event": "vision_analysis_error", "error": str(exc)}, ensure_ascii=False))
        return None


def _safe_name(value: str) -> str:
    return "".join(char if char.isalnum() or char in {"-", "_"} else "_" for char in value).strip("_") or "item"


def _resolve_material_blob(material: TaskMaterial, settings: Settings) -> MaterialBlob | None:
    if material.material_type == "drive_file" and material.drive_file_id:
        return _resolve_drive_material_blob(material, settings)
    if material.material_type == "link" and material.url:
        return _resolve_url_material_blob(material)
    return None


def _resolve_drive_material_blob(material: TaskMaterial, settings: Settings) -> MaterialBlob | None:
    try:
        from google.oauth2.credentials import Credentials
        from googleapiclient.discovery import build
        from googleapiclient.http import MediaIoBaseDownload
    except ImportError:
        return None

    if not settings.google_token_path.exists():
        return None

    try:
        creds = Credentials.from_authorized_user_file(str(settings.google_token_path), settings.google_scopes)
        service = build("drive", "v3", credentials=creds)
        metadata = service.files().get(fileId=material.drive_file_id, fields="name,mimeType").execute()
        mime_type = metadata.get("mimeType", "")
        filename = metadata.get("name") or material.title or material.material_id

        if mime_type in _GOOGLE_EXPORT_MIME_TYPES:
            export_mime_type = _GOOGLE_EXPORT_MIME_TYPES[mime_type]
            raw = service.files().export_media(
                fileId=material.drive_file_id,
                mimeType=export_mime_type,
            ).execute()
            suffix = ".pdf" if export_mime_type == "application/pdf" else ""
            return MaterialBlob(
                filename=_ensure_suffix(filename, suffix),
                mime_type=export_mime_type,
                data=raw if isinstance(raw, bytes) else bytes(raw),
            )

        if mime_type in _SUPPORTED_BINARY_MIME_TYPES or mime_type.startswith(_SUPPORTED_IMAGE_MIME_PREFIX):
            buf = io.BytesIO()
            downloader = MediaIoBaseDownload(buf, service.files().get_media(fileId=material.drive_file_id))
            done = False
            while not done:
                _, done = downloader.next_chunk()
            suffix = Path(filename).suffix or _suffix_for_mime_type(mime_type)
            return MaterialBlob(
                filename=_ensure_suffix(filename, suffix),
                mime_type=mime_type,
                data=buf.getvalue(),
            )
    except Exception as exc:
        print(
            json.dumps(
                {"event": "resolve_drive_blob_error", "drive_file_id": material.drive_file_id, "error": str(exc)},
                ensure_ascii=False,
            )
        )
        return None

    return None


def _resolve_url_material_blob(material: TaskMaterial) -> MaterialBlob | None:
    assert material.url is not None
    parsed = urlparse(material.url)
    suffix = Path(parsed.path).suffix.lower()
    if suffix not in _SUPPORTED_URL_SUFFIXES:
        return None

    try:
        with urlopen(material.url, timeout=30) as response:
            data = response.read()
            header_mime = response.headers.get_content_type()
    except Exception as exc:
        print(json.dumps({"event": "resolve_url_blob_error", "url": material.url, "error": str(exc)}, ensure_ascii=False))
        return None

    mime_type = header_mime or _mime_type_from_suffix(suffix)
    if mime_type not in _SUPPORTED_BINARY_MIME_TYPES and not mime_type.startswith(_SUPPORTED_IMAGE_MIME_PREFIX):
        mime_type = _mime_type_from_suffix(suffix)
    if mime_type not in _SUPPORTED_BINARY_MIME_TYPES and not mime_type.startswith(_SUPPORTED_IMAGE_MIME_PREFIX):
        return None

    filename = Path(parsed.path).name or f"{_safe_name(material.title)}{suffix}"
    return MaterialBlob(filename=filename, mime_type=mime_type, data=data)


def _analyze_with_document_intelligence(blob: MaterialBlob, settings: Settings) -> str | None:
    endpoint = settings.azure_document_intelligence_endpoint
    key = settings.azure_document_intelligence_key
    if not endpoint or not key:
        return None

    analyze_url = (
        endpoint.rstrip("/")
        + "/documentintelligence/documentModels/"
        + settings.azure_document_intelligence_model
        + ":analyze?"
        + urlencode(
            {
                "api-version": settings.azure_document_intelligence_api_version,
                "outputContentFormat": "markdown",
            }
        )
    )
    payload = json.dumps(
        {
            "base64Source": base64.b64encode(blob.data).decode("ascii"),
        }
    ).encode("utf-8")

    try:
        print(
            json.dumps(
                {
                    "event": "document_intelligence_started",
                    "filename": blob.filename,
                    "mime_type": blob.mime_type,
                    "size_bytes": len(blob.data),
                    "model": settings.azure_document_intelligence_model,
                    "api_version": settings.azure_document_intelligence_api_version,
                },
                ensure_ascii=False,
            )
        )
        request = Request(
            analyze_url,
            data=payload,
            headers={
                "Ocp-Apim-Subscription-Key": key,
                "Content-Type": "application/json",
            },
            method="POST",
        )
        with urlopen(request, timeout=60) as response:
            operation_location = response.headers.get("operation-location")
        if not operation_location:
            print(
                json.dumps(
                    {
                        "event": "document_intelligence_missing_operation_location",
                        "filename": blob.filename,
                    },
                    ensure_ascii=False,
                )
            )
            return None
        print(
            json.dumps(
                {
                    "event": "document_intelligence_accepted",
                    "filename": blob.filename,
                    "operation_location": operation_location,
                },
                ensure_ascii=False,
            )
        )

        for _ in range(30):
            poll_request = Request(
                operation_location,
                headers={"Ocp-Apim-Subscription-Key": key},
                method="GET",
            )
            with urlopen(poll_request, timeout=60) as poll_response:
                result = json.loads(poll_response.read().decode("utf-8"))

            status = (result.get("status") or "").lower()
            if status == "succeeded":
                analyze_result = result.get("analyzeResult") or {}
                content = analyze_result.get("content")
                if isinstance(content, str) and content.strip():
                    print(
                        json.dumps(
                            {
                                "event": "document_intelligence_succeeded",
                                "filename": blob.filename,
                                "content_length": len(content.strip()),
                            },
                            ensure_ascii=False,
                        )
                    )
                    return content.strip()
                print(
                    json.dumps(
                        {
                            "event": "document_intelligence_empty_result",
                            "filename": blob.filename,
                        },
                        ensure_ascii=False,
                    )
                )
                return None
            if status == "failed":
                print(
                    json.dumps(
                        {"event": "document_intelligence_failed", "filename": blob.filename, "result": result},
                        ensure_ascii=False,
                    )
                )
                return None
            time.sleep(1)
        print(
            json.dumps(
                {
                    "event": "document_intelligence_timeout",
                    "filename": blob.filename,
                    "poll_attempts": 30,
                },
                ensure_ascii=False,
            )
        )
    except Exception as exc:
        print(
            json.dumps(
                {"event": "document_intelligence_error", "filename": blob.filename, "error": str(exc)},
                ensure_ascii=False,
            )
        )
        return None

    return None


def _ensure_suffix(filename: str, suffix: str) -> str:
    if not suffix or filename.lower().endswith(suffix.lower()):
        return filename
    return filename + suffix


def _suffix_for_mime_type(mime_type: str) -> str:
    return {
        "application/pdf": ".pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/gif": ".gif",
        "image/webp": ".webp",
    }.get(mime_type, "")


def _mime_type_from_suffix(suffix: str) -> str:
    return {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif": "image/gif",
        ".webp": "image/webp",
    }.get(suffix.lower(), "application/octet-stream")
